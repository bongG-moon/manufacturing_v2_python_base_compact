from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from html import unescape
from pathlib import Path
from typing import Iterable

import pandas as pd
from openpyxl.styles import Alignment, Font, PatternFill

from summarize_100_case_results import CaseRow, load_cases


BASE_DIR = Path(__file__).resolve().parent
XLSX_PATH = BASE_DIR / "CASE_100_TEST_REPORT.xlsx"
DOCX_PATH = BASE_DIR / "CASE_100_TEST_REPORT.docx"


STATUS_SUCCESS = "м„ұкіө"
STATUS_AMBIGUOUS = "м• л§Ө"
STATUS_NEEDS_WORK = "к°ңм„ н•„мҡ”"

GREEN_FILL = PatternFill(fill_type="solid", fgColor="E2F0D9")
YELLOW_FILL = PatternFill(fill_type="solid", fgColor="FFF2CC")
RED_FILL = PatternFill(fill_type="solid", fgColor="F4CCCC")
HEADER_FILL = PatternFill(fill_type="solid", fgColor="D9EAF7")


@dataclass
class ExportRow:
    case_id: str
    scenario: str
    question: str
    tool: str
    status: str
    reason: str
    summary: str
    answer: str
    table_top5: str
    pandas_logic: str


def _clean_text(value: str) -> str:
    text = unescape(value or "")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _flatten_table_html(table_html: str) -> str:
    if not table_html:
        return ""

    if "<tr>" in table_html.lower():
        rows = re.findall(r"<tr>(.*?)</tr>", table_html, flags=re.IGNORECASE | re.DOTALL)
        flattened_rows: list[str] = []
        for row_html in rows:
            cells = re.findall(r"<t[dh]>(.*?)</t[dh]>", row_html, flags=re.IGNORECASE | re.DOTALL)
            cleaned = []
            for cell in cells:
                text = re.sub(r"<[^>]+>", " ", cell)
                text = _clean_text(text)
                if text:
                    cleaned.append(text)
            if cleaned:
                flattened_rows.append(" | ".join(cleaned))
        return "\n".join(flattened_rows)

    plain_text = _clean_text(table_html)
    if not plain_text:
        return ""

    chunks = re.split(r"(?=20\d{6})", plain_text)
    chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
    if len(chunks) <= 1:
        return plain_text

    header = chunks[0]
    data_rows = chunks[1:]
    formatted_rows = [header]
    formatted_rows.extend(data_rows[:5])
    return "\n".join(formatted_rows)


def _expected_columns_from_question(question: str) -> list[str]:
    mapping = {
        "лқјмқёлі„": ["лқјмқё"],
        "MODEлі„": ["MODE"],
        "modeлі„": ["MODE"],
        "DENлі„": ["DEN"],
        "denлі„": ["DEN"],
        "TECHлі„": ["TECH"],
        "techлі„": ["TECH"],
        "кіөм •кө°лі„": ["кіөм •кө°"],
        "кіөм •лі„": ["кіөм •"],
        "мЈјмҡ”л¶Ҳлҹүмң нҳ•лі„": ["мЈјмҡ”л¶Ҳлҹүмң нҳ•"],
        "л¶Ҳлҹүмң нҳ•лі„": ["мЈјмҡ”л¶Ҳлҹүмң нҳ•"],
    }
    expected: list[str] = []
    for token, columns in mapping.items():
        if token in question:
            expected.extend(columns)
    return expected


def _contains_error_text(text: str) -> bool:
    tokens = [
        "м§Җмӣҗн•ҳлҠ” мЎ°нҡҢ лҳҗлҠ” л¶„м„қ лІ”мң„лҘј лІ—м–ҙлӮ¬мҠөлӢҲлӢӨ",
        "м§Ҳл¬ё л¶„м„қ мӨ‘ л¬ём ңк°Җ л°ңмғқн–ҲмҠөлӢҲлӢӨ",
        "Traceback",
        "KeyError",
        "мӢӨнҢЁ",
        "мҳӨлҘҳ",
    ]
    return any(token in text for token in tokens)


def _classify_case(case: CaseRow) -> tuple[str, str]:
    answer = _clean_text(case.answer)
    summary = _clean_text(case.summary)
    table_text = _flatten_table_html(case.table_html)
    pandas_logic = _clean_text(case.pandas_logic)
    question = _clean_text(case.question)
    tool = _clean_text(case.tool)

    if not tool or tool == "None":
        return STATUS_NEEDS_WORK, "мӢӨн–үлҗң toolмқҙ м—Ҷм–ҙм„ң мӢӨм ң мІҳлҰ¬ кІҪлЎңлҘј нҷ•мқён•ҳкё° м–ҙл өмҠөлӢҲлӢӨ."

    if _contains_error_text(answer) or _contains_error_text(summary):
        return STATUS_NEEDS_WORK, "мқ‘лӢөмқҙлӮҳ мҡ”м•Ҫм—җ лӘ…мӢңм Ғмқё мҳӨлҘҳ/мӢӨнҢЁ л¬ёкө¬к°Җ нҸ¬н•Ёлҗҳм–ҙ мһҲмҠөлӢҲлӢӨ."

    if tool == "analyze_current_data" and not pandas_logic:
        return STATUS_NEEDS_WORK, "нӣ„мҶҚ л¶„м„қмқёлҚ° мғқм„ұлҗң pandas лЎңм§Ғмқҙ л№„м–ҙ мһҲмҠөлӢҲлӢӨ."

    expected_columns = _expected_columns_from_question(question)
    searchable = "\n".join([answer, table_text, pandas_logic])
    for expected in expected_columns:
        if expected not in searchable:
            return STATUS_AMBIGUOUS, f"м§Ҳл¬ёмқҖ `{expected}` кё°мӨҖ л¶„м„қмқ„ мҡ”кө¬н•ҳм§Җл§Ң кІ°кіј/лЎңм§Ғм—җм„ң к·ё кё°мӨҖмқҙ л¶„лӘ…н•ҳкІҢ л“ңлҹ¬лӮҳм§Җ м•ҠмҠөлӢҲлӢӨ."

    if "мғҒмң„" in question and tool == "analyze_current_data" and "head(" not in pandas_logic and "nlargest(" not in pandas_logic:
        return STATUS_AMBIGUOUS, "м§Ҳл¬ёмқҖ мғҒмң„ Nк°ңлҘј мҡ”кө¬н•ҳм§Җл§Ң pandas лЎңм§Ғм—җм„ң мғҒмң„ м¶”м¶ң мқҳлҸ„к°Җ лӘ…нҷ•н•ҳм§Җ м•ҠмҠөлӢҲлӢӨ."

    return STATUS_SUCCESS, "м§Ҳл¬ё мқҳлҸ„, мӢӨн–ү tool, кІ°кіј н…Ңмқҙлё”, pandas лЎңм§Ғмқҙ м „л°ҳм ҒмңјлЎң мһҗм—°мҠӨлҹҪкІҢ м—°кІ°лҗ©лӢҲлӢӨ."


def _build_export_rows() -> list[ExportRow]:
    export_rows: list[ExportRow] = []
    for case in load_cases():
        status, reason = _classify_case(case)
        export_rows.append(
            ExportRow(
                case_id=_clean_text(case.case_id),
                scenario=_clean_text(case.scenario),
                question=_clean_text(case.question),
                tool=_clean_text(case.tool),
                status=status,
                reason=reason,
                summary=_clean_text(case.summary),
                answer=_clean_text(case.answer),
                table_top5=_flatten_table_html(case.table_html),
                pandas_logic=_clean_text(case.pandas_logic),
            )
        )
    return export_rows


def _make_sheet_name(name: str, used_names: set[str]) -> str:
    base = re.sub(r"[:\\\\/?*\\[\\]]", "_", name).strip() or "Scenario"
    base = base[:31]
    candidate = base
    suffix = 2
    while candidate in used_names:
        suffix_text = f"_{suffix}"
        candidate = f"{base[:31 - len(suffix_text)]}{suffix_text}"
        suffix += 1
    used_names.add(candidate)
    return candidate


def _write_dataframe(writer: pd.ExcelWriter, sheet_name: str, df: pd.DataFrame) -> None:
    df.to_excel(writer, sheet_name=sheet_name, index=False)
    ws = writer.sheets[sheet_name]
    ws.auto_filter.ref = ws.dimensions
    ws.freeze_panes = "A2"

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = HEADER_FILL

    for column_cells in ws.columns:
        values = ["" if cell.value is None else str(cell.value) for cell in column_cells]
        width = min(max(len(max(values, key=len)) + 2, 12), 60)
        ws.column_dimensions[column_cells[0].column_letter].width = width


def _apply_status_colors(writer: pd.ExcelWriter, sheet_name: str, status_column_name: str = "мғҒнғң") -> None:
    ws = writer.sheets[sheet_name]
    status_col_idx = None
    for idx, cell in enumerate(ws[1], start=1):
        if cell.value == status_column_name:
            status_col_idx = idx
            break

    if status_col_idx is None:
        return

    for row in range(2, ws.max_row + 1):
        status_value = ws.cell(row=row, column=status_col_idx).value
        fill = None
        if status_value == STATUS_SUCCESS:
            fill = GREEN_FILL
        elif status_value == STATUS_AMBIGUOUS:
            fill = YELLOW_FILL
        elif status_value == STATUS_NEEDS_WORK:
            fill = RED_FILL

        if fill:
            for col in range(1, ws.max_column + 1):
                ws.cell(row=row, column=col).fill = fill


def _build_summary_dataframe(export_rows: Iterable[ExportRow]) -> pd.DataFrame:
    rows = list(export_rows)
    total = len(rows)
    success = sum(1 for row in rows if row.status == STATUS_SUCCESS)
    ambiguous = sum(1 for row in rows if row.status == STATUS_AMBIGUOUS)
    needs_work = sum(1 for row in rows if row.status == STATUS_NEEDS_WORK)

    summary_rows = [
        {"кө¬л¶„": "мҙқ мјҖмқҙмҠӨ", "к°’": total},
        {"кө¬л¶„": "м„ұкіө", "к°’": success},
        {"кө¬л¶„": "м• л§Ө", "к°’": ambiguous},
        {"кө¬л¶„": "к°ңм„ н•„мҡ”", "к°’": needs_work},
    ]

    return pd.DataFrame(summary_rows)


def _build_reason_dataframe(export_rows: Iterable[ExportRow]) -> pd.DataFrame:
    rows = [
        {
            "Case ID": row.case_id,
            "мғҒнғң": row.status,
            "мӢңлӮҳлҰ¬мҳӨ": row.scenario,
            "м§Ҳл¬ё": row.question,
            "мӮ¬мң ": row.reason,
        }
        for row in export_rows
        if row.status != STATUS_SUCCESS
    ]
    return pd.DataFrame(rows)


def _write_excel() -> None:
    export_rows = _build_export_rows()
    all_cases_df = pd.DataFrame(
        [
            {
                "Case ID": row.case_id,
                "мӢңлӮҳлҰ¬мҳӨ": row.scenario,
                "м§Ҳл¬ё": row.question,
                "Tool": row.tool,
                "мғҒнғң": row.status,
                "мӮ¬мң ": row.reason,
                "Summary": row.summary,
                "лӢөліҖ лӮҙмҡ©": row.answer,
                "н…Ңмқҙлё” мғҒмң„ 5к°ң": row.table_top5,
                "Pandas Logic": row.pandas_logic,
            }
            for row in export_rows
        ]
    )

    summary_df = _build_summary_dataframe(export_rows)
    review_df = _build_reason_dataframe(export_rows)

    with pd.ExcelWriter(XLSX_PATH, engine="openpyxl") as writer:
        _write_dataframe(writer, "100_CASES", all_cases_df)
        _apply_status_colors(writer, "100_CASES")

        _write_dataframe(writer, "SUMMARY", summary_df)
        _write_dataframe(writer, "QUALITY_REVIEW", review_df)
        _apply_status_colors(writer, "QUALITY_REVIEW")

        used_sheet_names = {"100_CASES", "SUMMARY", "QUALITY_REVIEW"}
        scenarios = sorted({row.scenario for row in export_rows})
        for scenario in scenarios:
            scenario_df = all_cases_df[all_cases_df["мӢңлӮҳлҰ¬мҳӨ"] == scenario].copy()
            sheet_name = _make_sheet_name(scenario, used_sheet_names)
            _write_dataframe(writer, sheet_name, scenario_df)
            _apply_status_colors(writer, sheet_name)


def _write_docx() -> bool:
    try:
        from docx import Document
    except Exception:
        return False

    export_rows = _build_export_rows()
    document = Document()
    document.add_heading("100 Case Test Report", level=1)
    document.add_paragraph("м§Ҳл¬ё, лӢөліҖ, мғҒмң„ 5к°ң н…Ңмқҙлё”, pandas лЎңм§Ғ, мғҒнғң нҢҗм •кіј мӮ¬мң лҘј н•ң л¬ём„ңм—җм„ң мқҪкё° мүҪкІҢ м •лҰ¬н•ң ліҙкі м„ңмһ…лӢҲлӢӨ.")

    document.add_heading("Summary", level=2)
    summary_df = _build_summary_dataframe(export_rows)
    for _, row in summary_df.iterrows():
        document.add_paragraph(f"{row['кө¬л¶„']}: {row['к°’']}")

    document.add_heading("Quality Review", level=2)
    review_df = _build_reason_dataframe(export_rows)
    if review_df.empty:
        document.add_paragraph("м• л§Ө лҳҗлҠ” к°ңм„ н•„мҡ” мјҖмқҙмҠӨк°Җ м—ҶмҠөлӢҲлӢӨ.")
    else:
        for _, row in review_df.iterrows():
            document.add_paragraph(f"Case {row['Case ID']} | {row['мғҒнғң']} | {row['м§Ҳл¬ё']}")
            document.add_paragraph(f"мӮ¬мң : {row['мӮ¬мң ']}")

    document.add_heading("Cases", level=2)
    for row in export_rows:
        document.add_heading(f"Case {row.case_id} - {row.scenario}", level=3)
        document.add_paragraph(f"м§Ҳл¬ё: {row.question}")
        document.add_paragraph(f"Tool: {row.tool}")
        document.add_paragraph(f"мғҒнғң: {row.status}")
        document.add_paragraph(f"мӮ¬мң : {row.reason}")
        document.add_paragraph(f"Summary: {row.summary}")
        document.add_paragraph(f"лӢөліҖ лӮҙмҡ©: {row.answer}")
        document.add_paragraph("н…Ңмқҙлё” мғҒмң„ 5к°ң:")
        document.add_paragraph(row.table_top5 or "кІ°кіј м—ҶмқҢ")
        document.add_paragraph("Pandas Logic:")
        document.add_paragraph(row.pandas_logic or "-")

    document.save(DOCX_PATH)
    return True


def main() -> None:
    _write_excel()
    docx_created = _write_docx()
    print(f"Saved Excel report to {XLSX_PATH}")
    if docx_created:
        print(f"Saved Word report to {DOCX_PATH}")
    else:
        print("Word report skipped: python-docx is not installed.")


if __name__ == "__main__":
    main()
